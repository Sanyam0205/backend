from pydantic import BaseModel, ValidationError
from openai import OpenAI
from dotenv import load_dotenv
import fitz  
import openai
from dotenv import load_dotenv
import os
load_dotenv()
import json
from Markdown2docx import Markdown2docx
import markdown
from docx import Document
from bs4 import BeautifulSoup
class profiletomarkdown:
    def __init__(self):
        openai_api_key = os.getenv("OPENAI_API_KEY")
        if not openai_api_key:
            raise ValueError("OpenAI API key is missing.")
        self.client = OpenAI(api_key=openai_api_key)

    def getmarkdown(self, profile: dict) -> dict:
        prompt = f"Convert the JSON {profile} to markdown. Do not temper with the content"

        try:
            completion = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a wonderful json to markdown converter"},
                    {"role": "user", "content": prompt},
                ]
            )
            print(completion.choices[0].message.content)
            event = completion.choices[0].message.content
            print(event)
            return event
        except Exception as e:
            print(f"Error with OpenAI: {e}")
        

    def markdown_profile_to_word(self,markdown_content, word_file):
        """
        Converts a Markdown string representing a researcher's profile to a Word (.docx) file.
        
        :param markdown_content: The Markdown content as a string (research profile).
        :param word_file: Path to the output Word (.docx) file.
        """
        html_content = markdown.markdown(markdown_content)
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'li', 'strong', 'em']):
            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'p':
                doc.add_paragraph(element.get_text())
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(f'• {li.get_text()}', style='ListBullet')
            elif element.name == 'strong':
                p = doc.add_paragraph()
                p.add_run(element.get_text()).bold = True
            elif element.name == 'em':
                p = doc.add_paragraph()
                p.add_run(element.get_text()).italic = True
        doc.save(word_file)
        print(f"Markdown content successfully saved to {word_file}")

    def runner(self,profile,word_path):
        res=self.getmarkdown(profile)
        self.markdown_profile_to_word(res,word_path)

# Example JSON data (as string)
profile_json = """
{
  "name": "Dr. John Doe",
  "title": "Researcher Profile",
  "research_areas": [
    "Machine Learning",
    "Natural Language Processing",
    "Computer Vision"
  ],
  "publications": [
    {
      "title": "Understanding Deep Learning",
      "author": "Doe, J.",
      "journal": "Journal of AI Research",
      "year": 2020
    },
    {
      "title": "Advances in NLP",
      "author": "Doe, J.",
      "journal": "International Journal of Computational Linguistics",
      "year": 2019
    },
    {
      "title": "A Survey on Machine Learning Models",
      "author": "Doe, J.",
      "journal": "Machine Intelligence Journal",
      "year": 2018
    }
  ],
  "awards": [
    "Best Paper Award, ICML 2020",
    "Outstanding Researcher Award, University of XYZ, 2019"
  ],
  "professional_memberships": [
    "Member, IEEE",
    "Member, Association for Computational Linguistics (ACL)"
  ]
}
"""
if __name__=="__main__":
    profile_data = json.loads(profile_json)
    x= profiletomarkdown()
    x.runner(profile_data,"output.docx")